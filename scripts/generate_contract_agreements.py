import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def style_contract_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="8" w:space="0" w:color="D4AF37"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="D4AF37"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_gold_divider(doc, thickness_pt=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="{int(thickness_pt * 8)}" w:space="1" w:color="D4AF37"/>
        </w:pBdr>
    ''')
    p._p.get_or_add_pPr().append(pBdr)

def add_section_banner(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="10" w:space="2" w:color="B8860B"/>
        </w:pBdr>
    ''')
    p._p.get_or_add_pPr().append(pBdr)

    r = p.add_run(title.upper())
    r.font.name = "Georgia"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(26, 28, 34)

def add_clause_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(f"CLAUSE {number}. ")
    r_num.font.name = "Arial"
    r_num.font.size = Pt(9.5)
    r_num.font.bold = True
    r_num.font.color.rgb = RGBColor(184, 134, 11) # Gold

    r_title = p.add_run(title.upper())
    r_title.font.name = "Georgia"
    r_title.font.size = Pt(9.5)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(26, 28, 34)

def add_body_p(doc, text, bold_prefix=None, space_after=4, is_justified=True):
    p = doc.add_paragraph()
    if is_justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(9.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(26, 28, 34)
        
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(40, 45, 55)
    return p

def save_document_safely(doc, output_path):
    try:
        doc.save(output_path)
        print(f"Successfully saved: {output_path}")
    except PermissionError:
        base, ext = os.path.splitext(output_path)
        alt_path = f"{base}_New{ext}"
        doc.save(alt_path)
        print(f"Warning: Original file was locked by Word. Saved to alternate path: {alt_path}")

def generate_indian_contract_law_agreement(branch_name, branch_address, jurisdiction_district, output_filename):
    doc = Document()
    
    # A4 Page Setup (0.55 in top/bottom, 0.65 in left/right)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "logo.png"))
    
    # ------------------ STAMP PAPER ADVISORY ------------------
    p_stamp_adv = doc.add_paragraph()
    p_stamp_adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_stamp_adv.paragraph_format.space_before = Pt(0)
    p_stamp_adv.paragraph_format.space_after = Pt(6)
    r_stamp = p_stamp_adv.add_run(
        "[ TO BE EXECUTED ON NON-JUDICIAL STAMP PAPER OF REQUISITE VALUE AS PRESCRIBED UNDER ARTICLE 5 OF THE KARNATAKA STAMP ACT, 1957 ]"
    )
    r_stamp.font.name = "Arial"
    r_stamp.font.size = Pt(7)
    r_stamp.font.bold = True
    r_stamp.font.color.rgb = RGBColor(120, 120, 120)

    # ------------------ HEADER WITH GSTIN ------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    header_table.columns[0].width = Inches(3.2)
    header_table.columns[1].width = Inches(3.77)

    # Left Cell: Logo
    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_left.paragraphs[0]
    p_logo.paragraph_format.space_after = Pt(0)
    
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.1))
    else:
        r_b = p_logo.add_run("DECOR8 INDIA")
        r_b.font.name = "Georgia"
        r_b.font.size = Pt(20)
        r_b.font.bold = True
        r_b.font.color.rgb = RGBColor(26, 28, 34)

    p_sub = cell_left.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(2)
    p_sub.paragraph_format.space_after = Pt(0)
    r_sub = p_sub.add_run(f"{branch_name.upper()} BRANCH • ARCHITECTURAL & TURNKEY STUDIO")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(7.5)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(184, 134, 11)

    # Right Cell: Entity & GSTIN Details
    cell_right = header_table.cell(0, 1)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)

    r_hq = p_right.add_run("DECOR8 INDIA INTERIORS PVT. LTD.\n")
    r_hq.font.name = "Georgia"
    r_hq.font.size = Pt(9.5)
    r_hq.font.bold = True
    r_hq.font.color.rgb = RGBColor(26, 28, 34)

    r_addr = p_right.add_run(
        "GSTIN: 29AABCD8826K1Z1 | PAN: AABCD8826K\n"
        f"{branch_address}\n"
        "Tel: +91 93805 23743 | Email: support@decor8india.com | Web: www.decor8india.com\n"
        "Registered Turnkey Architectural & Construction Contractor"
    )
    r_addr.font.name = "Arial"
    r_addr.font.size = Pt(7.2)
    r_addr.font.color.rgb = RGBColor(90, 95, 105)

    add_gold_divider(doc, thickness_pt=2.0)

    # ------------------ CONTRACT TITLE BANNER ------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(2)
    r_doc_title = p_title.add_run("TURNKEY ARCHITECTURAL & INTERIOR EXECUTION AGREEMENT")
    r_doc_title.font.name = "Georgia"
    r_doc_title.font.size = Pt(12)
    r_doc_title.font.bold = True
    r_doc_title.font.color.rgb = RGBColor(26, 28, 34)

    p_doc_sub = doc.add_paragraph()
    p_doc_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_doc_sub.paragraph_format.space_before = Pt(0)
    p_doc_sub.paragraph_format.space_after = Pt(10)
    r_doc_sub = p_doc_sub.add_run(f"[ Governed by The Indian Contract Act, 1872 • {branch_name} Branch Office Execution ]")
    r_doc_sub.font.name = "Arial"
    r_doc_sub.font.size = Pt(8.5)
    r_doc_sub.font.bold = True
    r_doc_sub.font.color.rgb = RGBColor(184, 134, 11)

    # ------------------ PREAMBLE & DESCRIPTION OF PARTIES ------------------
    add_body_p(
        doc,
        f"THIS TURNKEY ARCHITECTURAL & INTERIOR EXECUTION AGREEMENT (\"Agreement\") is made, entered into, and executed at "
        f"{jurisdiction_district}, State of Karnataka, India, on this _____ day of ________________________, 2026 (\"Execution Date\"), "
        f"BY AND BETWEEN:"
    )

    add_body_p(
        doc,
        f"M/S DECOR8 INDIA INTERIORS PVT. LTD., a company incorporated and governed under the provisions of the Companies Act, "
        f"holding GSTIN: 29AABCD8826K1Z1 and Permanent Account Number (PAN): AABCD8826K, having its operational branch office at "
        f"{branch_address}, acting through its authorized representative and Chief Executive Officer & Principal Architect, "
        f"MR. SATISH BHAT (hereinafter referred to as the \"FIRST PARTY\" or the \"CONTRACTOR / ARCHITECT\", which term shall, unless "
        f"repugnant to the context or meaning thereof, be deemed to mean and include its successors-in-interest, permitted assigns, "
        f"and administrators) OF THE FIRST PART;",
        bold_prefix="1. FIRST PARTY: "
    )

    p_and = doc.add_paragraph()
    p_and.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_and.paragraph_format.space_before = Pt(2)
    p_and.paragraph_format.space_after = Pt(2)
    r_and = p_and.add_run("— AND —")
    r_and.font.name = "Georgia"
    r_and.font.size = Pt(9)
    r_and.font.bold = True
    r_and.font.color.rgb = RGBColor(184, 134, 11)

    add_body_p(
        doc,
        "MR. / MRS. / M/S ____________________________________________________________________, son/daughter/wife of "
        "________________________________________________, aged about _____ years, holding Aadhaar No: ____________________________ "
        "and PAN: ____________________________, permanently residing at / having registered office at _________________________________"
        "____________________________________________________________________________________, Contact: +91 ________________________, "
        "Email: __________________________________________________ (hereinafter referred to as the \"SECOND PARTY\" or the "
        "\"CLIENT / PROPERTY OWNER\", which term shall, unless repugnant to the context or meaning thereof, be deemed to mean and "
        "include their legal heirs, executors, administrators, and permitted assigns) OF THE SECOND PART.",
        bold_prefix="2. SECOND PARTY: "
    )

    add_body_p(
        doc,
        "(The First Party and the Second Party are hereinafter collectively referred to as the \"Parties\" and individually as a \"Party\")."
    )

    # ------------------ RECITALS ------------------
    add_section_banner(doc, "RECITALS (WHEREAS CLAUSES)")

    add_body_p(
        doc,
        "A. WHEREAS the Second Party represents and warrants that they are the absolute, sole, and lawful owner / authorized title holder "
        "in physical possession of the immovable property situated at _______________________________________________________________"
        "____________________________________________________________________________________ (\"Project Site\"), admeasuring an "
        "estimated Super Built-up Area / Carpet Area of ______________ Sq. Ft."
    )

    add_body_p(
        doc,
        "B. WHEREAS the Second Party is desirous of carrying out complete architectural space planning, 2D/3D design visualization, "
        "structural modifications, custom modular woodwork, electrical, plumbing, false ceiling, surface finishing, and turnkey execution "
        "at the Project Site and has solicited the professional architectural services of the First Party."
    )

    add_body_p(
        doc,
        "C. WHEREAS the First Party possesses the necessary technical infrastructure, architectural expertise, qualified manpower, and "
        "manufacturing capabilities to execute turnkey residential and commercial interior projects as per recognized architectural standards."
    )

    add_body_p(
        doc,
        "D. WHEREAS the Parties have mutually agreed on the Scope of Work, Technical Material Specifications, Approved Bill of Quantities (BOQ), "
        "Commercial Considerations, and Milestone Payment Schedules as detailed in the Schedules attached hereto, and desire to formalize "
        "their mutual rights, covenants, and obligations under this Agreement in accordance with the provisions of The Indian Contract Act, 1872."
    )

    add_body_p(
        doc,
        "NOW, THEREFORE, IN CONSIDERATION OF THE MUTUAL COVENANTS, PROMISES, AND VALUABLE CONSIDERATION RECORDED HEREIN, THE ADEQUACY "
        "AND SUFFICIENCY OF WHICH IS HEREBY MUTUALLY ACKNOWLEDGED, THE PARTIES AGREE AS FOLLOWS:"
    )

    # ------------------ OPERATIVE CLAUSES ------------------
    add_section_banner(doc, "TERMS, COVENANTS & LEGAL CLAUSES")

    # Clause 1
    add_clause_heading(doc, "1", "Definitions & Interpretation")
    add_body_p(
        doc,
        "1.1. \"Turnkey Works\" shall mean the complete architectural planning, procurement, manufacturing, on-site installation, "
        "finishing, testing, and handover of interior infrastructure as specified in Schedule II."
    )
    add_body_p(
        doc,
        "1.2. \"Contract Consideration\" shall mean the total agreed financial value payable by the Second Party to the First Party for "
        "the due execution of Turnkey Works as specified in Clause 3."
    )
    add_body_p(
        doc,
        "1.3. \"Defect Liability Period (DLP)\" shall mean the period of 12 (twelve) calendar months commencing from the date of issuance "
        "of the Handover Certificate during which the Contractor shall rectify operational defects free of cost."
    )

    # Clause 2
    add_clause_heading(doc, "2", "Scope of Work & Architectural Obligations")
    add_body_p(
        doc,
        "2.1. The First Party shall execute the Turnkey Works strictly in accordance with approved architectural drawings, 3D renderings, "
        "and the Bill of Quantities (BOQ) set forth in Schedule II and Schedule III."
    )
    add_body_p(
        doc,
        "2.2. The turnkey scope encompasses: (a) 2D Furniture Layout Plans & 3D Visualizations; (b) Custom Modular Woodwork (Kitchens, Wardrobes, "
        "TV Units, Vanity Storage); (c) False Ceiling & Concealed Lighting; (d) Electrical Rewiring & Plumbing; (e) Wall Paneling, Marble/Tiling, "
        "and Premium Emulsion Painting; and (f) Post-installation Deep Cleaning & Joint Handover Inspection."
    )

    # Clause 3
    add_clause_heading(doc, "3", "Contract Price, GST & Financial Consideration")
    add_body_p(
        doc,
        "3.1. The total agreed Turnkey Contract Consideration for the complete execution of approved Scope/BOQ is fixed at ₹ ___________________ "
        "(Rupees ___________________________________________________________________________________ only) (\"Contract Price\")."
    )
    add_body_p(
        doc,
        "3.2. Goods & Services Tax (GST) shall be charged extra at the prevailing statutory rate (currently 18%) under GSTIN: 29AABCD8826K1Z1. "
        "All tax invoices issued by the First Party shall be compliant with the Central Goods and Services Tax Act, 2017."
    )

    # Clause 4
    add_clause_heading(doc, "4", "Milestone Payment Schedule & Realization Terms")
    add_body_p(
        doc,
        "4.1. The Second Party covenants to pay the Contract Price in strict accordance with the 4-Stage Milestone Schedule below, upon verification "
        "of physical completion of each designated stage:"
    )

    # Milestone Table
    m_table = doc.add_table(rows=5, cols=4)
    style_contract_table(m_table)
    m_table.columns[0].width = Inches(1.15)
    m_table.columns[1].width = Inches(0.95)
    m_table.columns[2].width = Inches(3.6)
    m_table.columns[3].width = Inches(1.27)

    headers = ["Milestone Stage", "% Share", "Milestone Trigger & Physical Deliverable", "Amount (₹) + GST"]
    for i, h in enumerate(headers):
        cell = m_table.cell(0, i)
        set_cell_background(cell, "1A1C22")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor(212, 175, 55)

    rows_data = [
        ("Milestone 1 (Token)", "10%", "Design Approval, 3D Renders, Detailed CAD Layouts, Site Measurement & Material Matrix sign-off", "₹ _____________"),
        ("Milestone 2 (Civil & Core)", "40%", "Core Procurement (IS:710 Marine Ply/Veneer), Civil alterations, Concealed Electrical & Plumbing start", "₹ _____________"),
        ("Milestone 3 (Fitout & Assembly)", "40%", "Delivery of Factory Modular Carcasses at Site, On-site Carpentry Fitout, False Ceiling, Wall Paneling & Base Paint", "₹ _____________"),
        ("Milestone 4 (Handover)", "10%", "Final Paint Coating, Hardware/Fixture Fitments, Deep Cleaning, Joint Snag Clearance & Formal Key Handover", "₹ _____________")
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        bg = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = m_table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if col_idx in (0, 1, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            if col_idx == 0:
                r.font.bold = True

    add_body_p(
        doc,
        "4.2. In the event of delay in milestone payment exceeding 7 (seven) days from invoice date, the First Party reserves the right to "
        "suspend site work, and interest at the rate of 12% per annum shall be chargeable on the outstanding sum under Section 73 of the Indian Contract Act.",
        space_after=4
    )

    # Clause 5
    add_clause_heading(doc, "5", "Project Timeline, Work Commencement & Completion")
    add_body_p(
        doc,
        "5.1. The agreed Turnkey Execution Timeline shall be ______________ (________________________) calendar working days, commencing "
        "from the date of realization of Milestone 2 payment and unhindered handover of physical site possession with necessary utility access."
    )
    add_body_p(
        doc,
        "5.2. Time shall be the essence of the contract subject to timely milestone clearances and material approvals by the Second Party."
    )

    # Clause 6
    add_clause_heading(doc, "6", "Material Specifications & Brand Quality Standards")
    add_body_p(
        doc,
        "6.1. All modular cabinetry carcasses and structural woodwork shall strictly conform to IS:710 Boiling Water Proof (BWP) Marine Grade "
        "or IS:303 Moisture Resistant (MR) Plywood specifications as detailed in Schedule III."
    )
    add_body_p(
        doc,
        "6.2. All hardware fittings, hinges, and drawer telescopic channels shall be original branded components (Hettich / Hafele / Ozone / Blum). "
        "Surface finishes shall utilize 1.0mm high-pressure laminates (Merino / Century / Greenlam) or certified natural wood veneers."
    )

    # Clause 7
    add_clause_heading(doc, "7", "Variations, Additions & Alterations (Variation Order Protocol)")
    add_body_p(
        doc,
        "7.1. No deviation, addition, or alteration to the approved drawings or BOQ shall be binding unless executed in writing via a mutual "
        "Variation Order (VO) specifying the cost differential and timeline extension."
    )

    # Clause 8
    add_clause_heading(doc, "8", "Client Obligations, Site Utilities & Statutory Permissions")
    add_body_p(
        doc,
        "8.1. The Second Party shall provide continuous, uninterrupted 3-phase electrical supply and potable water at the Project Site at their cost."
    )
    add_body_p(
        doc,
        "8.2. The Second Party shall secure all necessary Resident Welfare Association (RWA) permissions, society gate passes, and municipal NOCs."
    )

    # Clause 9
    add_clause_heading(doc, "9", "Labor Laws, Site Safety & Workmen Compliance")
    add_body_p(
        doc,
        "9.1. The First Party shall be solely responsible for the supervision, safety, and statutory compliance of its deployed workmen and subcontractors "
        "in accordance with the Workmen's Compensation Act, 1923, and applicable labor regulations."
    )

    # Clause 10
    add_clause_heading(doc, "10", "10-Year Structural Modular Warranty & Defect Liability")
    add_body_p(
        doc,
        "10.1. Decor8 India provides a 10-Year Structural Modular Warranty on factory-manufactured modular woodwork against delamination, "
        "borer/termite infestation, and manufacturing defects under normal residential usage."
    )
    add_body_p(
        doc,
        "10.2. During the 12-month Defect Liability Period (DLP), the First Party shall provide complimentary service and repair for movable hardware, "
        "hinges, and channels within 7 (seven) working days of written notification."
    )

    # Clause 11
    add_clause_heading(doc, "11", "Liquidated Damages & Genuine Pre-estimate of Delay")
    add_body_p(
        doc,
        "11.1. If the Contractor fails to complete the Turnkey Works within the agreed timeline (including justified extensions), the Contractor "
        "shall pay liquidated damages of 0.5% of the delayed contract value per week of delay, capped at a maximum of 5% under Section 74 of the Indian Contract Act."
    )

    # Clause 12
    add_clause_heading(doc, "12", "Force Majeure (Section 56, Indian Contract Act)")
    add_body_p(
        doc,
        "12.1. Neither Party shall be held liable for failure or delay in performance caused by Force Majeure events, including acts of God, "
        "floods, earthquakes, regional lockouts, government restrictions, pandemic measures, or civil unrest."
    )

    # Clause 13
    add_clause_heading(doc, "13", "Default & Termination Rights")
    add_body_p(
        doc,
        "13.1. Either Party may terminate this Agreement by giving 14 (fourteen) calendar days written notice in the event of a material breach "
        "not rectified within the notice window. Upon termination, a joint site audit shall determine the value of work executed and materials procured."
    )

    # Clause 14
    add_clause_heading(doc, "14", "Intellectual Property Rights")
    add_body_p(
        doc,
        "14.1. All architectural concepts, 3D renderings, CAD drawings, and custom modular engineering specifications developed by Decor8 India "
        "shall remain the intellectual property of the First Party. The Client is granted a non-exclusive license to use the designs solely for the Project Site."
    )

    # Clause 15
    add_clause_heading(doc, "15", "Indemnification & Limitation of Liability")
    add_body_p(
        doc,
        "15.1. Each Party agrees to indemnify and hold harmless the other Party against direct losses, third-party claims, or statutory penalties "
        "arising directly out of wilful misconduct or breach of statutory obligations under this Agreement."
    )

    # Clause 16
    add_clause_heading(doc, "16", "Dispute Resolution & Arbitration (Act of 1996)")
    add_body_p(
        doc,
        "16.1. In the event of any dispute or difference arising out of or in connection with this Agreement, the Parties shall first endeavor to settle "
        "the same amicably through mutual executive consultation within 15 (fifteen) days."
    )
    add_body_p(
        doc,
        f"16.2. If unresolved, the dispute shall be referred to and finally resolved by a Sole Arbitrator mutually appointed by both Parties in "
        f"accordance with the Arbitration and Conciliation Act, 1996. The seat and venue of arbitration shall be {jurisdiction_district}, Karnataka. "
        f"The arbitration proceedings shall be conducted in the English language."
    )

    # Clause 17
    add_clause_heading(doc, "17", "Governing Law & Exclusive Judicial Jurisdiction")
    add_body_p(
        doc,
        f"17.1. This Agreement shall be construed, interpreted, and governed in all respects in accordance with the substantive and procedural "
        f"Laws of the Republic of India. Subject to the arbitration clause, the competent Civil Courts situated at {jurisdiction_district}, "
        f"State of Karnataka, India, shall have exclusive jurisdiction over all matters arising hereunder."
    )

    # Clause 18
    add_clause_heading(doc, "18", "Entire Agreement, Severability & Amendments")
    add_body_p(
        doc,
        "18.1. This Agreement, along with Schedules I, II, III, and IV attached hereto, constitutes the entire understanding between the Parties "
        "and supersedes all prior oral or written negotiations, estimates, or representations."
    )

    # ------------------ TESTIMONIUM & EXECUTION BLOCK ------------------
    add_section_banner(doc, "TESTIMONIUM & EXECUTION IN WITNESS WHEREOF")

    add_body_p(
        doc,
        "IN WITNESS WHEREOF, the First Party and the Second Party hereto have set and subscribed their respective hands and seals "
        "to this Agreement on the day, month, and year first hereinabove written, in the presence of the following subscribing witnesses:"
    )

    sig_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Inches(3.5)
    sig_table.columns[1].width = Inches(3.47)

    # First Party Signature Box
    c_p1 = sig_table.cell(0, 0)
    p_p1 = c_p1.paragraphs[0]
    p_p1.paragraph_format.space_before = Pt(6)
    p_p1.paragraph_format.space_after = Pt(0)
    
    r1_h = p_p1.add_run("SIGNED, SEALED & DELIVERED BY THE FIRST PARTY:\n")
    r1_h.font.name = "Arial"
    r1_h.font.size = Pt(8.5)
    r1_h.font.bold = True
    r1_h.font.color.rgb = RGBColor(184, 134, 11)

    r1_txt = p_p1.add_run(
        "For DECOR8 INDIA INTERIORS PVT. LTD.\n"
        "GSTIN: 29AABCD8826K1Z1 | PAN: AABCD8826K\n\n\n\n"
        "_____________________________________________________\n"
        "MR. SATISH BHAT\n"
        "Chief Executive Officer & Principal Architect\n"
        f"Branch: {branch_name} Studio, Karnataka\n"
        "(Authorized Signatory / Official Seal)"
    )
    r1_txt.font.name = "Calibri"
    r1_txt.font.size = Pt(9)
    r1_txt.font.color.rgb = RGBColor(26, 28, 34)

    # Second Party Signature Box
    c_p2 = sig_table.cell(0, 1)
    p_p2 = c_p2.paragraphs[0]
    p_p2.paragraph_format.space_before = Pt(6)
    p_p2.paragraph_format.space_after = Pt(0)
    
    r2_h = p_p2.add_run("SIGNED, SEALED & DELIVERED BY THE SECOND PARTY:\n")
    r2_h.font.name = "Arial"
    r2_h.font.size = Pt(8.5)
    r2_h.font.bold = True
    r2_h.font.color.rgb = RGBColor(184, 134, 11)

    r2_txt = p_p2.add_run(
        "Client / Property Owner Signature:\n\n\n\n\n"
        "_____________________________________________________\n"
        "NAME: ______________________________________________\n"
        "Aadhaar No: _________________________________________\n"
        "PAN: ________________________________________________\n"
        "Contact No: +91 _____________________________________"
    )
    r2_txt.font.name = "Calibri"
    r2_txt.font.size = Pt(9)
    r2_txt.font.color.rgb = RGBColor(26, 28, 34)

    # Witnesses Box
    p_wit = doc.add_paragraph()
    p_wit.paragraph_format.space_before = Pt(16)
    p_wit.paragraph_format.space_after = Pt(4)
    p_wit.paragraph_format.keep_with_next = True
    r_wit_h = p_wit.add_run("IN THE PRESENCE OF INDEPENDENT WITNESSES:")
    r_wit_h.font.name = "Arial"
    r_wit_h.font.size = Pt(8.5)
    r_wit_h.font.bold = True
    r_wit_h.font.color.rgb = RGBColor(184, 134, 11)

    wit_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(wit_table)
    wit_table.columns[0].width = Inches(3.5)
    wit_table.columns[1].width = Inches(3.47)

    w1 = wit_table.cell(0, 0).paragraphs[0]
    w1.add_run(
        "WITNESS 1:\n"
        "Signature: __________________________________________\n"
        "Full Name: __________________________________________\n"
        "S/o / D/o: __________________________________________\n"
        "Address: ____________________________________________\n"
        "ID / Aadhaar No: ____________________________________"
    ).font.name = "Calibri"
    w1.runs[0].font.size = Pt(8.5)

    w2 = wit_table.cell(0, 1).paragraphs[0]
    w2.add_run(
        "WITNESS 2:\n"
        "Signature: __________________________________________\n"
        "Full Name: __________________________________________\n"
        "S/o / D/o: __________________________________________\n"
        "Address: ____________________________________________\n"
        "ID / Aadhaar No: ____________________________________"
    ).font.name = "Calibri"
    w2.runs[0].font.size = Pt(8.5)

    # ------------------ SCHEDULES & ANNEXURES ------------------
    add_section_banner(doc, "SCHEDULES & TECHNICAL ANNEXURES")

    add_body_p(
        doc,
        "SCHEDULE I: PROJECT SITE PARTICULARS\n"
        "• Project Name: ___________________________________________________\n"
        "• Site Address: ___________________________________________________\n"
        "• Super Built-up Area: ____________ Sq. Ft. | Carpet Area: ____________ Sq. Ft.\n"
        "• Configuration: [ ] 2BHK  [ ] 3BHK  [ ] 4BHK / Penthouse  [ ] Villa  [ ] Commercial Office",
        bold_prefix="SCHEDULE I: "
    )

    add_body_p(
        doc,
        "SCHEDULE II: SCOPE OF TURNKEY DELIVERABLES\n"
        "• 2D/3D Architectural Design & CAD layout sets with material palettes.\n"
        "• Foyer, Living & Dining: Shoe console, TV unit, wall paneling, false ceiling with warm LED coves.\n"
        "• Modular Kitchen: BWP Marine Grade carcass, soft-close tandem boxes, acrylic / PU shutters, quartz counter.\n"
        "• Master & Guest Bedrooms: Floor-to-ceiling wardrobes with loft, dresser unit, headboard acoustic paneling.\n"
        "• Electrical & Plumbing: Concealed wiring (Polycab/Havells), modular switches (Schneider/Legrand), CP fittings.\n"
        "• Surface Finishing: 2 coats putty, 1 coat primer, 2 coats Asian Paints Royale luxury emulsion.",
        bold_prefix="SCHEDULE II: "
    )

    add_body_p(
        doc,
        "SCHEDULE III: APPROVED MATERIAL GRADE MATRIX\n"
        "• Core Woodwork: IS:710 Boiling Water Proof (BWP) Marine Plywood (Century / Greenply / Kitply).\n"
        "• Surface Laminates: 1.0mm / 1.5mm High-Pressure Laminates (Merino / Greenlam / CenturyLaminates).\n"
        "• Hardware & Channels: Hettich Sensys 110° Soft-Close Hinges & InnoTech Atira Drawer Systems / Hafele.\n"
        "• Paints & Polishes: Asian Paints Royale Luxury Emulsion / PU Italian Wood Polish.",
        bold_prefix="SCHEDULE III: "
    )

    # Footer Line
    add_gold_divider(doc, thickness_pt=1.0)
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ftr.paragraph_format.space_before = Pt(2)
    p_ftr.paragraph_format.space_after = Pt(0)
    r_f = p_ftr.add_run(
        f"Decor8 India Turnkey Contract Agreement • GSTIN: 29AABCD8826K1Z1 • {branch_name} Branch • www.decor8india.com"
    )
    r_f.font.name = "Arial"
    r_f.font.size = Pt(7)
    r_f.font.bold = True
    r_f.font.color.rgb = RGBColor(140, 140, 140)

    save_document_safely(doc, output_filename)

def main():
    # 1. Bengaluru Branch Agreement
    bengaluru_addr = "#14, Sy No 36/1, Vasanth Vallabnagar, Vasanthpura, Uttrahalli Hobli, Bengaluru - 560061, Karnataka"
    bengaluru_out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Decor8India_Contract_Agreement_Bengaluru.docx"))
    generate_indian_contract_law_agreement("Bengaluru", bengaluru_addr, "Bengaluru Urban", bengaluru_out)

    # 2. Sirsi Branch Agreement
    sirsi_addr = "Decor8 India Design Studio, CP Bazaar / Hubli Road, Sirsi, Uttara Kannada - 581401, Karnataka"
    sirsi_out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Decor8India_Contract_Agreement_Sirsi.docx"))
    generate_indian_contract_law_agreement("Sirsi (Uttara Kannada)", sirsi_addr, "Sirsi, Uttara Kannada District", sirsi_out)

    # 3. Master Dual-Branch Unified Template
    master_addr = "Corporate HQ: Bengaluru (560061) & Regional Studio: Sirsi, Uttara Kannada (581401), Karnataka"
    master_out = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Decor8India_Turnkey_Contract_Agreement_Master.docx"))
    generate_indian_contract_law_agreement("Bengaluru & Sirsi", master_addr, "Bengaluru / Sirsi", master_out)

if __name__ == "__main__":
    main()
