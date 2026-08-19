import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tblBorders>
    ''')
    tblPr.append(tblBorders)

def add_gold_divider(doc, thickness_pt=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="{int(thickness_pt * 8)}" w:space="1" w:color="D4AF37"/>
        </w:pBdr>
    ''')
    p._p.get_or_add_pPr().append(pBdr)

def generate_letterhead():
    doc = Document()
    
    # Page setup - A4 with 0.6 inch margins for luxury professional feel
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "logo.png"))
    logo_icon_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "logo_icon.png"))
    
    # ------------------ HEADER SECTION ------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    
    # Column widths (Total ~ 6.97 inches)
    header_table.columns[0].width = Inches(3.2)
    header_table.columns[1].width = Inches(3.77)

    # Left Cell: Logo & Brand Name
    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_left.paragraphs[0]
    p_logo.paragraph_format.space_after = Pt(0)
    p_logo.paragraph_format.space_before = Pt(0)
    
    # Insert Logo if exists
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.2))
    elif os.path.exists(logo_icon_path):
        p_logo.add_run().add_picture(logo_icon_path, width=Inches(1.2))
    else:
        run_brand = p_logo.add_run("DECOR8 INDIA")
        run_brand.font.name = "Georgia"
        run_brand.font.size = Pt(22)
        run_brand.font.bold = True
        run_brand.font.color.rgb = RGBColor(26, 28, 34)

    # Subtitle under logo
    p_tag = cell_left.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(2)
    p_tag.paragraph_format.space_after = Pt(0)
    run_tag = p_tag.add_run("ARCHITECTURE & TURNKEY INTERIOR STUDIO")
    run_tag.font.name = "Arial"
    run_tag.font.size = Pt(7.5)
    run_tag.font.bold = True
    run_tag.font.color.rgb = RGBColor(184, 134, 11) # Dark Gold

    # Right Cell: Registration & Official Contact Header
    cell_right = header_table.cell(0, 1)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(1)
    p_right.paragraph_format.space_before = Pt(0)

    run_hq = p_right.add_run("DECOR8 INDIA INTERIORS PVT. LTD.\n")
    run_hq.font.name = "Georgia"
    run_hq.font.size = Pt(9.5)
    run_hq.font.bold = True
    run_hq.font.color.rgb = RGBColor(26, 28, 34)

    run_details = p_right.add_run(
        "GSTIN: 29AABCD8826K1Z1 | PAN: AABCD8826K\n"
        "Corporate HQ: #14, Sy No 36/1, Vasanth Vallabnagar, Bengaluru - 560061\n"
        "Tel: +91 93805 23743 | +91 98765 43210 | support@decor8india.com\n"
        "Web: www.decor8india.com | ISO 9001:2015 Certified Turnkey Studio"
    )
    run_details.font.name = "Arial"
    run_details.font.size = Pt(7.5)
    run_details.font.color.rgb = RGBColor(90, 95, 105)

    # Gold Accent Line under Header
    add_gold_divider(doc, thickness_pt=2.0)

    # ------------------ METADATA BAR (Ref, Date) ------------------
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta_table)
    meta_table.columns[0].width = Inches(3.5)
    meta_table.columns[1].width = Inches(3.47)

    # Ref No
    p_ref = meta_table.cell(0, 0).paragraphs[0]
    p_ref.paragraph_format.space_before = Pt(4)
    p_ref.paragraph_format.space_after = Pt(4)
    run_ref_lbl = p_ref.add_run("REF NO: ")
    run_ref_lbl.font.name = "Arial"
    run_ref_lbl.font.size = Pt(8.5)
    run_ref_lbl.font.bold = True
    run_ref_lbl.font.color.rgb = RGBColor(184, 134, 11)

    run_ref_val = p_ref.add_run("D8I / 2026-27 / _________")
    run_ref_val.font.name = "Consolas"
    run_ref_val.font.size = Pt(8.5)
    run_ref_val.font.color.rgb = RGBColor(50, 50, 50)

    # Date
    p_date = meta_table.cell(0, 1).paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(4)
    p_date.paragraph_format.space_after = Pt(4)
    run_date_lbl = p_date.add_run("DATE: ")
    run_date_lbl.font.name = "Arial"
    run_date_lbl.font.size = Pt(8.5)
    run_date_lbl.font.bold = True
    run_date_lbl.font.color.rgb = RGBColor(184, 134, 11)

    run_date_val = p_date.add_run("____ / ____ / 2026")
    run_date_val.font.name = "Consolas"
    run_date_val.font.size = Pt(8.5)
    run_date_val.font.color.rgb = RGBColor(50, 50, 50)

    # Thin separator under meta
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(14)
    pBdr_thin = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="4" w:space="1" w:color="E0E0E0"/>
        </w:pBdr>
    ''')
    p_sep._p.get_or_add_pPr().append(pBdr_thin)

    # ------------------ RECIPIENT & SUBJECT ------------------
    p_to = doc.add_paragraph()
    p_to.paragraph_format.space_before = Pt(0)
    p_to.paragraph_format.space_after = Pt(4)
    r_to_lbl = p_to.add_run("TO,\n")
    r_to_lbl.font.name = "Arial"
    r_to_lbl.font.size = Pt(9.5)
    r_to_lbl.font.bold = True
    r_to_lbl.font.color.rgb = RGBColor(26, 28, 34)

    r_to_val = p_to.add_run(
        "[Client / Organization Name]\n"
        "[Designation / Department]\n"
        "[Address / City, State, PIN Code]\n"
        "Contact: +91 [Mobile Number] | Email: [Client Email]"
    )
    r_to_val.font.name = "Calibri"
    r_to_val.font.size = Pt(9.5)
    r_to_val.font.color.rgb = RGBColor(70, 75, 85)

    # Subject Line
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(8)
    p_sub.paragraph_format.space_after = Pt(12)
    p_sub_border = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:left w:val="single" w:sz="16" w:space="6" w:color="D4AF37"/>
        </w:pBdr>
    ''')
    p_sub._p.get_or_add_pPr().append(p_sub_border)
    
    r_sub_lbl = p_sub.add_run("SUBJECT: ")
    r_sub_lbl.font.name = "Arial"
    r_sub_lbl.font.size = Pt(10)
    r_sub_lbl.font.bold = True
    r_sub_lbl.font.color.rgb = RGBColor(184, 134, 11)

    r_sub_txt = p_sub.add_run("Official Letter / Architectural Scope & Proposal Confirmation")
    r_sub_txt.font.name = "Georgia"
    r_sub_txt.font.size = Pt(10)
    r_sub_txt.font.bold = True
    r_sub_txt.font.color.rgb = RGBColor(26, 28, 34)

    # Salutation
    p_sal = doc.add_paragraph()
    p_sal.paragraph_format.space_before = Pt(4)
    p_sal.paragraph_format.space_after = Pt(8)
    r_sal = p_sal.add_run("Dear Sir / Madam,")
    r_sal.font.name = "Calibri"
    r_sal.font.size = Pt(10)
    r_sal.font.bold = True

    # Body Paragraph 1
    p_body1 = doc.add_paragraph()
    p_body1.paragraph_format.space_before = Pt(0)
    p_body1.paragraph_format.space_after = Pt(8)
    p_body1.paragraph_format.line_spacing = 1.15
    r_b1 = p_body1.add_run(
        "We are pleased to present this official communication on behalf of Decor8 India Architecture & Turnkey Interior Studio. "
        "With a dedicated focus on architectural excellence, premium material curation, and bespoke craftsmanship, our studio "
        "is committed to delivering world-class residential and commercial spaces across India."
    )
    r_b1.font.name = "Calibri"
    r_b1.font.size = Pt(10)
    r_b1.font.color.rgb = RGBColor(40, 45, 55)

    # Body Paragraph 2 (Placeholder for actual letter content)
    p_body2 = doc.add_paragraph()
    p_body2.paragraph_format.space_before = Pt(0)
    p_body2.paragraph_format.space_after = Pt(12)
    p_body2.paragraph_format.line_spacing = 1.15
    r_b2 = p_body2.add_run(
        "[Type your official letter body, project milestones, contractual obligations, scope of architectural deliverables, "
        "or quotation terms here. You can freely edit this document in Microsoft Word while maintaining the exact corporate branding, "
        "typography, and footer layout.]"
    )
    r_b2.font.name = "Calibri"
    r_b2.font.size = Pt(10)
    r_b2.font.italic = True
    r_b2.font.color.rgb = RGBColor(100, 105, 115)

    # ------------------ SIGNATURE BLOCK ------------------
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    sig_table.columns[0].width = Inches(4.0)
    sig_table.columns[1].width = Inches(2.97)

    # Left: Warm Regards & Authorized Signatory
    cell_sig_left = sig_table.cell(0, 0)
    p_sig = cell_sig_left.paragraphs[0]
    p_sig.paragraph_format.space_before = Pt(14)
    p_sig.paragraph_format.space_after = Pt(2)
    r_warm = p_sig.add_run("Yours Faithfully,\n")
    r_warm.font.name = "Calibri"
    r_warm.font.size = Pt(9.5)

    r_comp = p_sig.add_run("For DECOR8 INDIA INTERIORS PVT. LTD.\n\n\n")
    r_comp.font.name = "Arial"
    r_comp.font.size = Pt(9)
    r_comp.font.bold = True
    r_comp.font.color.rgb = RGBColor(26, 28, 34)

    r_name = p_sig.add_run("MR. SATISH BHAT\n")
    r_name.font.name = "Georgia"
    r_name.font.size = Pt(10)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(184, 134, 11)

    r_desig = p_sig.add_run("Chief Executive Officer & Principal Architect\nCouncil of Architecture Reg. | Karnataka, India")
    r_desig.font.name = "Arial"
    r_desig.font.size = Pt(8)
    r_desig.font.color.rgb = RGBColor(90, 95, 105)

    # Right: Stamp & Signature Box
    cell_sig_right = sig_table.cell(0, 1)
    cell_sig_right.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    p_stamp = cell_sig_right.paragraphs[0]
    p_stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_stamp.paragraph_format.space_before = Pt(20)
    
    r_seal = p_stamp.add_run("[ Official Seal & Signature ]\n")
    r_seal.font.name = "Arial"
    r_seal.font.size = Pt(8)
    r_seal.font.italic = True
    r_seal.font.color.rgb = RGBColor(150, 150, 150)

    # ------------------ FOOTER SECTION ------------------
    # Divider above footer
    p_ftr_div = doc.add_paragraph()
    p_ftr_div.paragraph_format.space_before = Pt(24)
    p_ftr_div.paragraph_format.space_after = Pt(4)
    pBdr_ftr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="12" w:space="1" w:color="D4AF37"/>
        </w:pBdr>
    ''')
    p_ftr_div._p.get_or_add_pPr().append(pBdr_ftr)

    # 3-Column Branch Offices Table
    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(footer_table)
    
    col_w = Inches(2.32)
    footer_table.columns[0].width = col_w
    footer_table.columns[1].width = col_w
    footer_table.columns[2].width = col_w

    # Branch 1: Bengaluru
    c1 = footer_table.cell(0, 0)
    p_c1 = c1.paragraphs[0]
    p_c1.paragraph_format.space_before = Pt(2)
    p_c1.paragraph_format.space_after = Pt(0)
    r1_h = p_c1.add_run("BENGALURU (HQ Studio)\n")
    r1_h.font.name = "Arial"
    r1_h.font.size = Pt(7.5)
    r1_h.font.bold = True
    r1_h.font.color.rgb = RGBColor(184, 134, 11)
    
    r1_b = p_c1.add_run("#14, Sy No 36/1, Vasanth Vallabnagar,\nUttrahalli Hobli, Bengaluru - 560061\nTel: +91 93805 23743")
    r1_b.font.name = "Arial"
    r1_b.font.size = Pt(6.8)
    r1_b.font.color.rgb = RGBColor(100, 105, 115)

    # Branch 2: Hyderabad
    c2 = footer_table.cell(0, 1)
    p_c2 = c2.paragraphs[0]
    p_c2.paragraph_format.space_before = Pt(2)
    p_c2.paragraph_format.space_after = Pt(0)
    r2_h = p_c2.add_run("HYDERABAD Studio\n")
    r2_h.font.name = "Arial"
    r2_h.font.size = Pt(7.5)
    r2_h.font.bold = True
    r2_h.font.color.rgb = RGBColor(184, 134, 11)
    
    r2_b = p_c2.add_run("Plot 450, Road 36, Jubilee Hills,\nHyderabad, Telangana - 500033\nTel: +91 98765 43210")
    r2_b.font.name = "Arial"
    r2_b.font.size = Pt(6.8)
    r2_b.font.color.rgb = RGBColor(100, 105, 115)

    # Branch 3: Mumbai
    c3 = footer_table.cell(0, 2)
    p_c3 = c3.paragraphs[0]
    p_c3.paragraph_format.space_before = Pt(2)
    p_c3.paragraph_format.space_after = Pt(0)
    r3_h = p_c3.add_run("MUMBAI Studio\n")
    r3_h.font.name = "Arial"
    r3_h.font.size = Pt(7.5)
    r3_h.font.bold = True
    r3_h.font.color.rgb = RGBColor(184, 134, 11)
    
    r3_b = p_c3.add_run("Suite 802, Maker Chambers V,\nNariman Point, Mumbai - 400021\nTel: +91 98200 11223")
    r3_b.font.name = "Arial"
    r3_b.font.size = Pt(6.8)
    r3_b.font.color.rgb = RGBColor(100, 105, 115)

    # Final Bottom Bar
    p_bottom = doc.add_paragraph()
    p_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bottom.paragraph_format.space_before = Pt(6)
    p_bottom.paragraph_format.space_after = Pt(0)
    r_bot = p_bottom.add_run("www.decor8india.com  •  support@decor8india.com  •  Luxury Interior Architecture & Turnkey Execution")
    r_bot.font.name = "Arial"
    r_bot.font.size = Pt(7)
    r_bot.font.bold = True
    r_bot.font.color.rgb = RGBColor(140, 140, 140)

def save_document_safely(doc, output_path):
    try:
        doc.save(output_path)
        print(f"Successfully generated: {output_path}")
    except PermissionError:
        base, ext = os.path.splitext(output_path)
        alt_path = f"{base}_New{ext}"
        doc.save(alt_path)
        print(f"Original file was open in Word. Saved to: {alt_path}")

    # Save Sample Word document
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Decor8India_Official_Letterhead.docx"))
    save_document_safely(doc, output_path)

    # Generate Blank Version
    generate_blank_letterhead()

def generate_blank_letterhead():
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "logo.png"))
    logo_icon_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "logo_icon.png"))
    
    # ------------------ HEADER SECTION ------------------
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)
    header_table.columns[0].width = Inches(3.2)
    header_table.columns[1].width = Inches(3.77)

    cell_left = header_table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_left.paragraphs[0]
    p_logo.paragraph_format.space_after = Pt(0)
    p_logo.paragraph_format.space_before = Pt(0)
    
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(2.2))
    elif os.path.exists(logo_icon_path):
        p_logo.add_run().add_picture(logo_icon_path, width=Inches(1.2))

    p_tag = cell_left.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(2)
    p_tag.paragraph_format.space_after = Pt(0)
    run_tag = p_tag.add_run("ARCHITECTURE & TURNKEY INTERIOR STUDIO")
    run_tag.font.name = "Arial"
    run_tag.font.size = Pt(7.5)
    run_tag.font.bold = True
    run_tag.font.color.rgb = RGBColor(184, 134, 11)

    cell_right = header_table.cell(0, 1)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(1)
    p_right.paragraph_format.space_before = Pt(0)

    run_hq = p_right.add_run("DECOR8 INDIA INTERIORS PVT. LTD.\n")
    run_hq.font.name = "Georgia"
    run_hq.font.size = Pt(9.5)
    run_hq.font.bold = True
    run_hq.font.color.rgb = RGBColor(26, 28, 34)

    run_details = p_right.add_run(
        "GSTIN: 29AABCD8826K1Z1 | PAN: AABCD8826K\n"
        "Corporate HQ: #14, Sy No 36/1, Vasanth Vallabnagar, Bengaluru - 560061\n"
        "Tel: +91 93805 23743 | +91 98765 43210 | support@decor8india.com\n"
        "Web: www.decor8india.com | ISO 9001:2015 Certified Turnkey Studio"
    )
    run_details.font.name = "Arial"
    run_details.font.size = Pt(7.5)
    run_details.font.color.rgb = RGBColor(90, 95, 105)

    add_gold_divider(doc, thickness_pt=2.0)

    # Blank Lines for Typing
    for _ in range(12):
        p_blank = doc.add_paragraph()
        p_blank.paragraph_format.space_before = Pt(12)
        p_blank.paragraph_format.space_after = Pt(12)

    # ------------------ FOOTER SECTION ------------------
    p_ftr_div = doc.add_paragraph()
    p_ftr_div.paragraph_format.space_before = Pt(24)
    p_ftr_div.paragraph_format.space_after = Pt(4)
    pBdr_ftr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="12" w:space="1" w:color="D4AF37"/>
        </w:pBdr>
    ''')
    p_ftr_div._p.get_or_add_pPr().append(pBdr_ftr)

    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(footer_table)
    
    col_w = Inches(2.32)
    footer_table.columns[0].width = col_w
    footer_table.columns[1].width = col_w
    footer_table.columns[2].width = col_w

    c1 = footer_table.cell(0, 0)
    p_c1 = c1.paragraphs[0]
    p_c1.paragraph_format.space_before = Pt(2)
    p_c1.paragraph_format.space_after = Pt(0)
    r1_h = p_c1.add_run("BENGALURU (HQ Studio)\n")
    r1_h.font.name = "Arial"
    r1_h.font.size = Pt(7.5)
    r1_h.font.bold = True
    r1_h.font.color.rgb = RGBColor(184, 134, 11)
    r1_b = p_c1.add_run("#14, Sy No 36/1, Vasanth Vallabnagar,\nUttrahalli Hobli, Bengaluru - 560061\nTel: +91 93805 23743")
    r1_b.font.name = "Arial"
    r1_b.font.size = Pt(6.8)
    r1_b.font.color.rgb = RGBColor(100, 105, 115)

    c2 = footer_table.cell(0, 1)
    p_c2 = c2.paragraphs[0]
    p_c2.paragraph_format.space_before = Pt(2)
    p_c2.paragraph_format.space_after = Pt(0)
    r2_h = p_c2.add_run("HYDERABAD Studio\n")
    r2_h.font.name = "Arial"
    r2_h.font.size = Pt(7.5)
    r2_h.font.bold = True
    r2_h.font.color.rgb = RGBColor(184, 134, 11)
    r2_b = p_c2.add_run("Plot 450, Road 36, Jubilee Hills,\nHyderabad, Telangana - 500033\nTel: +91 98765 43210")
    r2_b.font.name = "Arial"
    r2_b.font.size = Pt(6.8)
    r2_b.font.color.rgb = RGBColor(100, 105, 115)

    c3 = footer_table.cell(0, 2)
    p_c3 = c3.paragraphs[0]
    p_c3.paragraph_format.space_before = Pt(2)
    p_c3.paragraph_format.space_after = Pt(0)
    r3_h = p_c3.add_run("MUMBAI Studio\n")
    r3_h.font.name = "Arial"
    r3_h.font.size = Pt(7.5)
    r3_h.font.bold = True
    r3_h.font.color.rgb = RGBColor(184, 134, 11)
    r3_b = p_c3.add_run("Suite 802, Maker Chambers V,\nNariman Point, Mumbai - 400021\nTel: +91 98200 11223")
    r3_b.font.name = "Arial"
    r3_b.font.size = Pt(6.8)
    r3_b.font.color.rgb = RGBColor(100, 105, 115)

    p_bottom = doc.add_paragraph()
    p_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bottom.paragraph_format.space_before = Pt(6)
    p_bottom.paragraph_format.space_after = Pt(0)
    r_bot = p_bottom.add_run("www.decor8india.com  •  support@decor8india.com  •  Luxury Interior Architecture & Turnkey Execution")
    r_bot.font.name = "Arial"
    r_bot.font.size = Pt(7)
    r_bot.font.bold = True
    r_bot.font.color.rgb = RGBColor(140, 140, 140)

    blank_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Decor8India_Letterhead_Blank.docx"))
    save_document_safely(doc, blank_path)

if __name__ == "__main__":
    generate_letterhead()
